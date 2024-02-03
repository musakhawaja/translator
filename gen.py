from google.cloud import documentai_v1 as documentai
from google.oauth2 import service_account
import io
from PyPDF2 import PdfFileReader, PdfFileWriter
from openai import OpenAI
import json
from docx import Document
import base64
import re
import math
import tempfile
from pydub import AudioSegment
import os
from dotenv import load_dotenv
import concurrent.futures


load_dotenv()
client = OpenAI(api_key = os.getenv('OPENAI_API_KEY'))

def format_timestamp(seconds):
    """Helper function to format timestamps."""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    seconds = seconds % 60
    return f"{int(hours):02}:{int(minutes):02}:{int(seconds):02}"

def read_docx(file_path):
    doc = Document(file_path)
    full_text = ""
    char_count = 0

    for para in doc.paragraphs:
        para_text = para.text + "\n"
        full_text += para_text
        char_count += len(para_text)

        if char_count >= 3000:
            full_text += "\n--EndOfPage--\n\n"
            char_count = 0  
    return full_text.strip()

def audio_transcript(audio_file):
    audio = AudioSegment.from_file(audio_file)
    length_audio = len(audio) / 1000  # Convert to seconds
    full_transcription = ""
    if length_audio <= 60:
        start_seconds = 0 * 60
        end_seconds = min((0 + 1) * 60, length_audio)
        start = 0 * 60 * 1000  # Convert to milliseconds for slicing
        end = min((0 + 1) * 60 * 1000, len(audio))
        transcription = client.audio.transcriptions.create(model="whisper-1", file=audio_file, response_format="text")
        # full_transcription = transcription.strip()
        timestamp = f"Start:[{format_timestamp(start_seconds)}] End:[{format_timestamp(end_seconds)}]"
        full_transcription += f"{timestamp}\n{transcription.strip()}\n\n--EndOfPage--\n\n"
    else:
        for i in range(0, math.ceil(length_audio / 60)):
            start_seconds = i * 60
            end_seconds = min((i + 1) * 60, length_audio)
            start = i * 60 * 1000  # Convert to milliseconds for slicing
            end = min((i + 1) * 60 * 1000, len(audio))
            chunk = audio[start:end]
            with tempfile.NamedTemporaryFile(suffix='.wav', delete=True) as temp_file:
                chunk.export(temp_file.name, format="wav")
                transcription = client.audio.transcriptions.create(model="whisper-1", file=open(temp_file.name, 'rb'), response_format="text")
                timestamp = f"Start:[{format_timestamp(start_seconds)}] End:[{format_timestamp(end_seconds)}]"
                full_transcription += f"{timestamp}\n{transcription.strip()}\n\n--EndOfPage--\n\n"

    print(full_transcription)
    return full_transcription

def split_pdf_to_chunks(uploaded_file, pages_per_chunk=15):
    file_stream = io.BytesIO(uploaded_file.getvalue())
    reader = PdfFileReader(file_stream)
    total_pages = reader.getNumPages()

    for start_page in range(0, total_pages, pages_per_chunk):
        writer = PdfFileWriter()
        end_page = min(start_page + pages_per_chunk, total_pages)
        for page_number in range(start_page, end_page):
            writer.addPage(reader.getPage(page_number))
        chunk_file = io.BytesIO()
        writer.write(chunk_file)
        chunk_file.seek(0)  
        yield chunk_file

def read_document(chunk):
    credentials = service_account.Credentials.from_service_account_file("ocrproject-412113-82a31889338f.json")
    client = documentai.DocumentProcessorServiceClient(credentials=credentials)
    name = "projects/707177808576/locations/us/processors/6eebcb4a15b88393"

    content = chunk.read()
    raw_document = documentai.RawDocument(content=content, mime_type="application/pdf")
    request = documentai.ProcessRequest(name=name, raw_document=raw_document)
    result = client.process_document(request=request)
    document = result.document

    structured_text = ""
    for page in document.pages:
        rows = {}
        for block in page.blocks:
            y_coord = block.layout.bounding_poly.vertices[0].y
            if y_coord not in rows:
                rows[y_coord] = []
            block_text = document.text[block.layout.text_anchor.text_segments[0].start_index:
                                       block.layout.text_anchor.text_segments[0].end_index]
            rows[y_coord].append((block.layout.bounding_poly.vertices[0].x, block_text))

        for y_coord in sorted(rows.keys()):
            row = sorted(rows[y_coord], key=lambda x: x[0])  
            row_text = '\t'.join([text for _, text in row])  
            structured_text += row_text + "\n"

        structured_text += "--EndOfPage--\n\n"

    return structured_text

def translate(file_path, prompt, source_lang="English", target_lang="Urdu", model_version="gpt-4-turbo"):
    with open(file_path, 'r', encoding='utf-8') as file:
        text = file.read()

    # Model selection based on input parameter
    model = "gpt-4-0125-preview" if model_version == "gpt-4-turbo" else "gpt-4"

    try:
        completion = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": f"Translate from {source_lang} to {target_lang}: {prompt}"},
                {"role": "user", "content": text}
            ]
        )

        result = completion.choices[0].message.content
        print(result)
        return ("success", result)  # Success
    except Exception as exc:
        if "rate limit" in str(exc).lower():  # Check for rate limit
            if model_version == "gpt-4-turbo":
                return ("retry_with_gpt4", None)  # Retry with GPT-4
            else:
                return ("requeue", None)  # Requeue for later processing
        else:
            raise

def translate_and_combine_text(edited_text, prompt, source_lang, target_lang):
    pages = edited_text.split("--EndOfPage--")
    temp_file_paths = []
    indexed_translated_texts = []  # Store translations with their original index

    # Save each page to a temp file with index
    for index, page in enumerate(pages):
        with tempfile.NamedTemporaryFile(delete=False, mode='w+', encoding='utf-8', suffix=".txt") as temp_file:
            temp_file.write(page)
            temp_file.flush()  # Make sure data is written to disk
            temp_file_paths.append((temp_file.name, index))  # Store path with index

    # Translate each temp file using multiprocessing, preserving index
    with concurrent.futures.ThreadPoolExecutor(max_workers=20) as executor:
        # Initial submission with GPT-4 Turbo
        future_to_file = {
            executor.submit(translate, file_path, prompt, source_lang, target_lang, "gpt-4-turbo"): (file_path, index, "gpt-4-turbo")
            for file_path, index in temp_file_paths
        }
        
        to_requeue = []

        while future_to_file:
            done, _ = concurrent.futures.wait(future_to_file, return_when=concurrent.futures.FIRST_COMPLETED)
            for future in done:
                file_path, index, model_version = future_to_file.pop(future)
                result, translated_text = future.result()
                
                if result == "success":
                    print(f"Page {index} translated successfully")
                    indexed_translated_texts.append((index, translated_text))
                elif result == "retry_with_gpt4":
                    print(f"Page {index} failed. Retrying with GPT4 Base")
                    retry_future = executor.submit(translate, file_path, prompt, source_lang, target_lang, "gpt-4")
                    future_to_file[retry_future] = (file_path, index, "gpt-4")
                elif result == "requeue":
                    print(f"Page {index} failed. REQUE")
                    to_requeue.append((file_path, index))
                else:
                    raise ValueError(f"Unexpected result from translate: {result}")

        # Requeue tasks that hit rate limits on both models
        for file_path, index in to_requeue:
            print(f"Requeuing {file_path} due to persistent rate limits...")
            requeue_future = executor.submit(translate, file_path, prompt, source_lang, target_lang, "gpt-4-turbo")
            future_to_file[requeue_future] = (file_path, index, "gpt-4-turbo")


    # Sort translated texts by their index and then combine
    indexed_translated_texts.sort(key=lambda x: x[0])  # Sort by index
    combined_translated_text = "\n\n".join([text for _, text in indexed_translated_texts])

    # Cleanup: delete temp files
    for file_path, _ in temp_file_paths:
        os.remove(file_path)

    return combined_translated_text


# def translate_and_combine_text(edited_text, prompt, source_lang, target_lang):
#     pages = edited_text.split("--EndOfPage--")
#     temp_file_paths = []
#     translated_texts = []

#     # Save each page to a temp file
#     for page in pages:
#         with tempfile.NamedTemporaryFile(delete=False, mode='w+', encoding='utf-8', suffix=".txt") as temp_file:
#             temp_file.write(page)
#             temp_file.flush()  # Make sure data is written to disk
#             temp_file_paths.append(temp_file.name)

#     # Translate each temp file
#     for file_path in temp_file_paths:
#         translated_text = translate(file_path, prompt, source_lang, target_lang)
#         translated_texts.append(translated_text)

#     # Combine translated texts
#     combined_translated_text = "\n\n".join(translated_texts)

#     # Cleanup: delete temp files
#     for file_path in temp_file_paths:
#         os.remove(file_path)

#     return combined_translated_text

def clean_text(text):
    """
    Removes characters that are not compatible with XML (e.g., NULL bytes, control characters)
    except for tab (\t), newline (\n), and carriage return (\r).
    """
    # Allow only printable characters and specific control characters (\t, \n, \r)
    return ''.join(char for char in text if char.isprintable() or char in '\t\n\r')

def convert_text_to_docx_bytes(text):
    doc = Document()
    lines = text.split('\n')
    for line in lines:
        # Clean line to remove invalid XML characters
        cleaned_line = clean_text(line)
        paragraph = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', cleaned_line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    docx_io = io.BytesIO()
    doc.save(docx_io)
    docx_io.seek(0)
    return docx_io
